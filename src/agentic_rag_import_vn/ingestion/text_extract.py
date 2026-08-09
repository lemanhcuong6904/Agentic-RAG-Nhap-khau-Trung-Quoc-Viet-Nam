from __future__ import annotations

import json
import logging
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable

import pandas as pd
from tqdm import tqdm

from agentic_rag_import_vn.config import settings
from agentic_rag_import_vn.io_utils import ensure_dirs, read_table, write_table
from agentic_rag_import_vn.quality.audit import finish_run, new_run, write_events, write_processing_run
from agentic_rag_import_vn.quality.text import text_quality


LEGAL_DOCUMENT_ROLES = {"legal_general", "origin_legal", "vat_legal"}
TABLE_STRATEGIES = {"spreadsheet_table", "structured_table"}

logging.getLogger("pypdf").setLevel(logging.ERROR)


def page_record(
    doc: pd.Series,
    page: int | None,
    section: str | None,
    text: str,
    source: str,
    blocks: list[dict[str, object]] | None = None,
    tables: list[dict[str, object]] | None = None,
) -> dict[str, object]:
    quality = text_quality(text)
    status = "pass" if quality["status"] == "pass" and doc.get("risk_level") != "high" else "review"
    if "[PAGE_EXTRACTION_ERROR]" in text:
        status = "fail"
        quality["needs_review"] = True
        quality["reasons"] = [*quality.get("reasons", []), "page_extraction_error"]
    return {
        "document_id": doc["document_id"],
        "title": doc["title"],
        "relative_path": doc["relative_path"],
        "category": doc["category"],
        "document_role": doc["document_role"],
        "agreement": doc.get("agreement"),
        "page": page,
        "section": section,
        "text_source": source,
        "parser": doc.get("parser"),
        "parser_version": settings.parser_version,
        "blocks": blocks or [{"block_id": "b1", "type": "paragraph", "text": text, "bbox": None, "confidence": None}],
        "tables": tables or [],
        "quality": {**quality, "status": status, "quality_status": status},
        "created_at": datetime.now(timezone.utc).isoformat(),
    }


def extract_pdf_pages(path: Path, doc: pd.Series) -> list[dict[str, object]]:
    from pypdf import PdfReader

    pages: list[dict[str, object]] = []
    reader = PdfReader(str(path), strict=False)
    for index, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception as exc:
            text = f"[PAGE_EXTRACTION_ERROR] {exc!r}"
        pages.append(page_record(doc, index, None, text.strip(), "pdf_text"))
    return pages


def extract_docx_pages(path: Path, doc: pd.Series) -> list[dict[str, object]]:
    from docx import Document

    document = Document(str(path))
    blocks: list[dict[str, object]] = []
    text_parts: list[str] = []
    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text.strip()
        if not text:
            continue
        style = paragraph.style.name.lower() if paragraph.style else ""
        block_type = "heading" if "heading" in style or text.lower().startswith(("điều ", "dieu ")) else "paragraph"
        blocks.append({"block_id": f"p{index}", "type": block_type, "text": text, "bbox": None, "confidence": None})
        text_parts.append(text)
    tables: list[dict[str, object]] = []
    for table_index, table in enumerate(document.tables, start=1):
        table_rows = []
        for row_index, row in enumerate(table.rows, start=1):
            values = [cell.text.strip() for cell in row.cells]
            table_rows.append(values)
            if any(values):
                text_parts.append(" | ".join(values))
        tables.append({"table_id": f"t{table_index}", "rows": table_rows})
    return [page_record(doc, None, None, "\n".join(text_parts).strip(), "docx", blocks=blocks, tables=tables)]


def extract_spreadsheet_tables(path: Path, doc: pd.Series) -> tuple[list[dict[str, object]], list[dict[str, object]]]:
    if path.suffix.lower() == ".csv":
        sheets = {"csv": pd.read_csv(path, dtype=str, header=None)}
    else:
        sheets = pd.read_excel(path, sheet_name=None, dtype=str, header=None)

    pages: list[dict[str, object]] = []
    cells: list[dict[str, object]] = []
    for sheet_name, df in sheets.items():
        df = df.dropna(how="all")
        lines: list[str] = []
        for row_index, row in df.iterrows():
            values = []
            for col_index, value in enumerate(row.tolist()):
                if pd.isna(value) or not str(value).strip():
                    continue
                text_value = str(value).strip()
                values.append(text_value)
                cells.append(
                    {
                        "source_document_id": doc["document_id"],
                        "source_path": doc["relative_path"],
                        "sheet": str(sheet_name),
                        "row_number": int(row_index) + 1,
                        "column_number": int(col_index) + 1,
                        "raw_value": text_value,
                        "normalized_value": text_value,
                        "quality_status": "review" if doc.get("risk_level") == "high" else "pass",
                    }
                )
            if values:
                lines.append(" | ".join(values))
        pages.append(page_record(doc, None, str(sheet_name), "\n".join(lines), "spreadsheet_table"))
    return pages, cells


def extract_document(path: Path, doc: pd.Series) -> tuple[list[dict[str, object]], list[dict[str, object]]]:
    file_type = str(doc["file_type"])
    if file_type == "pdf":
        return extract_pdf_pages(path, doc), []
    if file_type == "docx":
        return extract_docx_pages(path, doc), []
    if file_type in {"xlsx", "xls", "csv"}:
        return extract_spreadsheet_tables(path, doc)
    if file_type == "txt":
        text = path.read_text(encoding="utf-8", errors="ignore")
        return [page_record(doc, None, None, text, "text_file")], []
    raise ValueError(f"Unsupported extraction type: {file_type}")


def markdown_front_matter(doc: pd.Series) -> str:
    fields = {
        "document_id": doc["document_id"],
        "title": str(doc["title"]).replace('"', "'"),
        "category": doc["category"],
        "document_role": doc["document_role"],
        "agreement": doc.get("agreement") or "",
        "language": doc.get("language") or "",
        "source_file": doc["relative_path"],
        "source_url": doc.get("source_url") or "",
        "effective_from": doc.get("effective_from") or "",
        "effective_to": doc.get("effective_to") or "",
        "status": doc.get("status") or "unknown",
        "parser_version": settings.parser_version,
    }
    lines = ["---"]
    for key, value in fields.items():
        lines.append(f"{key}: \"{value}\"")
    lines.append("---")
    return "\n".join(lines)


def write_markdown(doc: pd.Series, pages: Iterable[dict[str, object]]) -> Path:
    target = settings.interim_dir / "markdown" / f"{doc['document_id']}.md"
    target.parent.mkdir(parents=True, exist_ok=True)
    body = [markdown_front_matter(doc), "", f"# {doc['title']}", ""]
    for page in pages:
        if page.get("page"):
            body.append(f"<!-- page: {page['page']} -->")
        elif page.get("section"):
            body.append(f"<!-- section: {page['section']} -->")
        body.append("")
        for block in page.get("blocks", []):
            text = str(block.get("text") or "").strip()
            if not text:
                continue
            if block.get("type") == "heading":
                body.append(f"## {text}")
            else:
                body.append(text)
            body.append("")
    target.write_text("\n".join(body).strip() + "\n", encoding="utf-8")
    return target


def run_text_extraction(limit: int | None = None) -> Path:
    run = new_run("extract_text", {"limit": limit})
    ensure_dirs(
        [
            settings.interim_dir / "parsed_json",
            settings.interim_dir / "markdown",
            settings.interim_dir / "raw_tables",
            settings.interim_dir / "page_text",
            settings.reports_dir / "extraction_quality",
            settings.manifests_dir,
        ]
    )
    registry = read_table(settings.processed_dir / "document_registry.parquet")
    candidates = registry[
        registry["document_role"].isin(LEGAL_DOCUMENT_ROLES)
        & registry["file_type"].isin({"pdf", "docx", "txt"})
        & registry["duplicate_of"].isna()
    ].copy()
    table_candidates = registry[
        registry["parse_strategy"].isin(TABLE_STRATEGIES)
        & registry["file_type"].isin({"xlsx", "xls", "csv"})
        & registry["duplicate_of"].isna()
    ].copy()
    candidates = pd.concat([candidates, table_candidates], ignore_index=True)
    if limit:
        candidates = candidates.head(limit)

    parsed_path = settings.interim_dir / "parsed_json" / "pages.jsonl"
    legacy_page_text_path = settings.interim_dir / "page_text" / "documents.jsonl"
    errors: list[dict[str, object]] = []
    quality_rows: list[dict[str, object]] = []
    table_cells: list[dict[str, object]] = []
    markdown_paths: list[Path] = []

    with parsed_path.open("w", encoding="utf-8") as parsed_out, legacy_page_text_path.open("w", encoding="utf-8") as legacy_out:
        for _, doc in tqdm(candidates.iterrows(), total=len(candidates), desc="parse-documents"):
            file_path = settings.project_root / str(doc["relative_path"])
            try:
                pages, cells = extract_document(file_path, doc)
            except Exception as exc:
                errors.append(
                    {
                        "document_id": doc["document_id"],
                        "relative_path": doc["relative_path"],
                        "stage": "extract_text",
                        "error": repr(exc),
                        "created_at": datetime.now(timezone.utc).isoformat(),
                    }
                )
                continue
            table_cells.extend(cells)
            if doc["document_role"] in LEGAL_DOCUMENT_ROLES:
                markdown_paths.append(write_markdown(doc, pages))
            for page in pages:
                parsed_out.write(json.dumps(page, ensure_ascii=False) + "\n")
                legacy_out.write(
                    json.dumps(
                        {
                            "document_id": page["document_id"],
                            "title": page["title"],
                            "relative_path": page["relative_path"],
                            "category": page["category"],
                            "agreement": page.get("agreement"),
                            "page": page.get("page"),
                            "section": page.get("section"),
                            "quality_status": page["quality"]["quality_status"],
                            "text": " ".join(str(block.get("text") or "") for block in page.get("blocks", [])),
                        },
                        ensure_ascii=False,
                    )
                    + "\n"
                )
                quality_rows.append(
                    {
                        "document_id": page["document_id"],
                        "relative_path": page["relative_path"],
                        "page": page.get("page"),
                        "section": page.get("section"),
                        "document_role": page["document_role"],
                        **page["quality"],
                    }
                )

    if errors:
        write_table(pd.DataFrame(errors), settings.manifests_dir / "ingestion_errors.parquet")
    else:
        stale_errors = settings.manifests_dir / "ingestion_errors.parquet"
        if stale_errors.exists():
            stale_errors.unlink()
    if table_cells:
        write_table(pd.DataFrame(table_cells), settings.interim_dir / "raw_tables" / "cells.parquet")
    quality_path = write_table(pd.DataFrame(quality_rows), settings.reports_dir / "extraction_quality" / "page_quality.parquet")
    write_processing_run(
        finish_run(
            run,
            metrics={
                "documents_attempted": len(candidates),
                "pages_written": len(quality_rows),
                "markdown_files": len(markdown_paths),
                "errors": len(errors),
            },
        )
    )
    write_events(
        [
            {
                "run_id": run["run_id"],
                "stage": "extract_text",
                "event_type": "artifact_written",
                "artifact": str(parsed_path),
                "created_at": datetime.now(timezone.utc).isoformat(),
            },
            {
                "run_id": run["run_id"],
                "stage": "extract_text",
                "event_type": "quality_report_written",
                "artifact": str(quality_path),
                "created_at": datetime.now(timezone.utc).isoformat(),
            },
        ]
    )
    return parsed_path
